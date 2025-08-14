import re
import requests
import docx
import io
import json
import openai
from typing import List, Dict, Optional
from jinja2 import Template
import json_repair
from chatbot.llm_models.llm_script import handle_bedrock_model


def extract_tags_from_document_sections(url: str, openai_api_key: Optional[str] = None,
                                       company_bot=None) -> Dict[str, List[str]]:
    """
    Extract tags/classification directly from document sections with AI fallbacks
    
    Args:
        url: Document URL (Google Docs, DOCX, etc.)
        openai_api_key: OpenAI API key for GPT-3.5 Turbo fallback
        company_bot: Company bot object for Bedrock model fallback
    
    Returns:
        Dictionary with extracted tags and classifications (cleaned)
    """
    
    def get_document_paragraphs(url: str) -> List[str]:
        """Get all paragraphs from document"""
        
        def get_google_doc_text(doc_url: str) -> str:
            match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', doc_url)
            if not match:
                raise ValueError("Invalid Google Docs URL")
            
            doc_id = match.group(1)
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
            response = requests.get(export_url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Split into paragraphs
            return [p.strip() for p in response.text.split('\n') if p.strip()]
        
        def get_docx_paragraphs(content: bytes) -> List[str]:
            doc = docx.Document(io.BytesIO(content))
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            return paragraphs
        
        try:
            if 'docs.google.com/document' in url:
                return get_google_doc_text(url)
            
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            if url.lower().endswith(('.docx', '.doc')):
                return get_docx_paragraphs(response.content)
            else:
                # Treat as text
                return [p.strip() for p in response.text.split('\n') if p.strip()]
                
        except Exception as e:
            print(f"Error reading document: {e}")
            return []
    
    def get_full_document_text(url: str) -> str:
        """Get complete document text for AI processing"""
        try:
            if 'docs.google.com/document' in url:
                match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', url)
                if not match:
                    raise ValueError("Invalid Google Docs URL")
                
                doc_id = match.group(1)
                export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
                
                headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
                response = requests.get(export_url, headers=headers, timeout=30)
                response.raise_for_status()
                return response.text
            
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            if url.lower().endswith(('.docx', '.doc')):
                doc = docx.Document(io.BytesIO(response.content))
                text_parts = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text)
                return '\n'.join(text_parts)
            else:
                return response.text
                
        except Exception as e:
            print(f"Error getting full document text: {e}")
            return ""

    def find_section_content(paragraphs: List[str], section_keywords: List[str]) -> List[str]:
        """Find content under specific section headings"""

        content = []
        found_section = False

        for i, para in enumerate(paragraphs):
            para_lower = para.lower().strip()

            # Check if this paragraph is a section heading we're looking for
            is_section_header = False
            for keyword in section_keywords:
                if (keyword.lower() in para_lower and
                        len(para.split()) <= 5 and  # Short heading
                        (para_lower.startswith(keyword.lower()) or
                         para_lower.endswith(keyword.lower()) or
                         para_lower == keyword.lower())):
                    is_section_header = True
                    found_section = True
                    break

            if is_section_header:
                continue  # Skip the header itself

            # If we found our section, collect content until next major heading
            if found_section:
                # Check if this is a numbered section header (e.g., "12. Next Section")
                numbered_section = re.match(r'^\d+\.\s+\w+', para)

                # Only stop if we hit a numbered section that looks like a major heading
                if numbered_section and len(para) > 15:  # Longer numbered sections are likely new sections
                    break

                # Or stop if line ends with colon and has section keywords (but not list items)
                if (para.endswith(':') and
                        not re.match(r'^\s*[\-\*•\[\]x\s]', para) and
                        any(keyword in para_lower for keyword in
                            ['section', 'chapter', 'part', 'instructions', 'purpose',
                             'overview', 'background', 'audience', 'intended users',
                             'complementary resources', 'limitations', 'alignment'])):
                    break

                # Add content from this section
                content.append(para)

        return content

    def parse_list_content(content_lines: List[str]) -> List[str]:
        """Parse content lines into clean list items - FIXED VERSION"""

        # Define all possible checkbox/checkmark indicators
        CHECKBOX_MARKERS = [
            '[x]', '[X]', '[ ]', '✅', '☑️', '✓', '✔', '☒', '☐', '⬜', '🔲', '🔳',
        ]

        # Build regex pattern from markers
        # Escape special regex characters in markers
        escaped_markers = [re.escape(marker) for marker in CHECKBOX_MARKERS]
        markers_pattern = '|'.join(escaped_markers)

        items = []

        for line in content_lines:
            line = line.strip()
            if not line:
                continue

            # Check if line contains any of our markers
            if any(marker in line for marker in CHECKBOX_MARKERS):
                # Find all checkbox patterns in the line
                # Pattern matches: optional bullet + optional space + (any marker) + space + (capture everything until another marker or end)
                # Using negative lookahead to stop before next marker
                checkbox_pattern = rf'[\-\*•]?\s*(?:{markers_pattern})\s*([^✅☑️✓✔☒☐⬜🔲🔳\[\]]+?)(?=[\-\*•]?\s*(?:{markers_pattern})|$)'
                matches = re.findall(checkbox_pattern, line, re.IGNORECASE)
                for match in matches:
                    item = match.strip()
                    if item:
                        items.append(item)
                if matches:  # If we found matches, continue to next line
                    continue

            # Handle other list formats...

            # Format 1: Bullet points: - item, • item, * item
            bullet_match = re.match(r'^\s*[\-\*•]\s*(.+)', line)
            if bullet_match:
                item = bullet_match.group(1).strip()
                # Remove any markers from the item
                item = re.sub(rf'(?:{markers_pattern})\s*', '', item, flags=re.IGNORECASE).strip()
                if item:
                    items.append(item)
                continue

            # Format 2: Numbered lists: 1. item, 1) item
            number_match = re.match(r'^\s*\d+[.)]\s*(.+)', line)
            if number_match:
                item = number_match.group(1).strip()
                # Remove any markers from the item
                item = re.sub(rf'(?:{markers_pattern})\s*', '', item, flags=re.IGNORECASE).strip()
                if item:
                    items.append(item)
                continue

            # Format 3: Comma/semicolon separated in single line
            has_markers = any(marker in line for marker in CHECKBOX_MARKERS)
            if (',' in line or ';' in line) and not has_markers and not any(marker in line for marker in ['-', '*']):
                sub_items = re.split(r'[,;]', line)
                for sub_item in sub_items:
                    clean_item = sub_item.strip()
                    if clean_item:
                        items.append(clean_item)
                continue

            # Format 4: Plain text item (if it doesn't look like a heading)
            if len(line.split()) <= 6 and line and not line.endswith(':'):
                # Remove any markers
                clean_line = re.sub(rf'(?:{markers_pattern})\s*', '', line, flags=re.IGNORECASE).strip()
                if clean_line:
                    items.append(clean_line)

        # Final cleanup of items
        cleaned_items = []
        for item in items:
            # Remove extra whitespace
            clean_item = re.sub(r'\s+', ' ', item).strip()
            # Remove leading/trailing punctuation
            clean_item = clean_item.strip('.,;:-')
            # Remove any remaining markers
            clean_item = re.sub(rf'(?:{markers_pattern})', '', clean_item, flags=re.IGNORECASE).strip()

            if clean_item and len(clean_item) > 1:
                cleaned_items.append(clean_item)

        return cleaned_items

    def extract_tags_with_openai(document_text: str, api_key: str) -> Dict[str, List[str]]:
        """
        Fallback 1: Extract tags using OpenAI GPT-3.5 Turbo
        
        Args:
            document_text: Full document content
            api_key: OpenAI API key
        
        Returns:
            Dictionary with extracted tags and classifications
        """
        try:
            print("Attempting OpenAI GPT-3.5 Turbo extraction...")
            
            # Initialize OpenAI client
            client = openai.OpenAI(api_key=api_key)
            
            # Truncate text if too long
            if len(document_text) > 8000:
                document_text = document_text[:8000] + "..."
            
            prompt = f"""
Analyze this document and extract tags and classification information.

Document Content:
{document_text}

Look for existing sections labeled "Tags", "Classification", "Type", "Category" or similar.
If such sections exist, extract the items from those lists.
If no explicit sections exist, generate relevant tags based on the content.

Return ONLY valid JSON in this exact format:
{{
  "tags": ["tag1", "tag2", "tag3"],
  "classification": ["category1", "category2"]
}}

Focus on:
- Educational content and terminology
- Government and policy-related terms
- M&E (Monitoring & Evaluation) concepts
- Field work and assessment terms
- Tools and resources
- Administrative content

Extract 5-10 relevant tags maximum.
"""
            
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert document analyzer. Extract tags and classification information from documents in JSON format."
                    },
                    {"role": "user", "content": prompt}
                ],
                max_tokens=800,
                temperature=0.3
            )
            
            result_text = response.choices[0].message.content.strip()
            
            # Parse JSON response
            try:
                result = json.loads(result_text)
                print(f"OpenAI extraction successful: {result}")
                return {
                    'tags': result.get('tags', []),
                    'classification': result.get('classification', [])
                }
            except json.JSONDecodeError:
                # Fallback parsing
                print("JSON parsing failed, attempting regex extraction...")
                tags = re.findall(r'"([^"]*)"', result_text)
                return {
                    'tags': tags[:10] if tags else [],
                    'classification': []
                }
                
        except Exception as e:
            print(f"OpenAI extraction failed: {str(e)}")
            return {'tags': [], 'classification': []}
    
    def extract_tags_with_bedrock(document_text: str, company_bot) -> Dict[str, List[str]]:
        try:
            print("Attempting Bedrock model extraction...")
            
            if len(document_text) > 8000:
                document_text = document_text[:8000] + "..."
            
            system_prompt = [
                {
                    'text': company_bot.context
                },
            ]
            tag_context = company_bot.tag_context
            if not tag_context:
                return {'tags': [], 'classification': []}

            context_data = {
                "document_text": document_text,
            }
            template = Template(tag_context)
            tag_context = template.render(context_data)

            messages = [{
                'role': 'user',
                'content': [{'text': f"{tag_context}"}]
            }]

            tool = company_bot.tool_context
            if tool and isinstance(tool, str):
                tool = json_repair.repair_json(tool, return_objects=True)
            
            response = handle_bedrock_model(
                system_prompt=system_prompt, messages=messages, model_name=company_bot.llm_model,
                temperature=company_bot.bot_temperature, max_token=company_bot.max_token, company_bot=company_bot,
                tools=tool
            )

            print("response: ", response)
            print("type: response: ", type(response))

            if response and isinstance(response, dict):
                extracted_data = response.pop("parameters", response.pop("input", None))
                if extracted_data and isinstance(extracted_data, dict):
                    response.clear()
                    response.update(extracted_data)
                print("last response: ", response)
                print("last type: response: ", type(response))
                return response
            else:
                result = {'tags': [], 'classification': []}
            
            print(f"Bedrock extraction successful: {result}")
            return {
                'tags': result.get('tags', []),
                'classification': result.get('classification', [])
            }
            
        except Exception as e:
            print(f"Bedrock extraction failed: {str(e)}")
            return {'tags': [], 'classification': []}
    
    # Main extraction logic with fallbacks
    try:
        print(f"Reading document: {url}")
        
        # Get all paragraphs from document
        paragraphs = get_document_paragraphs(url)
        print(f"Found {len(paragraphs)} paragraphs")
        
        result = {}
        manual_extraction_successful = False
        
        # Try manual extraction first
        try:
            # Look for Tags section (including "Tags / Classification")
            tag_keywords = ['tags', 'tag', 'keywords', 'labels', 'tags / classification', 'classification']
            tag_content = find_section_content(paragraphs, tag_keywords)
            if tag_content:
                result['tags'] = parse_list_content(tag_content)
                print(f"Found tags section with {len(result['tags'])} items: {result['tags']}")
                manual_extraction_successful = True
            
            # Look for Classification section separately
            classification_keywords = ['classification', 'category', 'categories', 'type', 'types']
            classification_content = find_section_content(paragraphs, classification_keywords)
            if classification_content:
                result['classification'] = parse_list_content(classification_content)
                print(f"Found classification section with {len(result['classification'])} items: {result['classification']}")
                manual_extraction_successful = True

            if (not result.get('tags')) and (not result.get('classification')):
                manual_extraction_successful = False

        except Exception as e:
            print(f"Manual extraction failed: {e}")
            manual_extraction_successful = False
        
        # If manual extraction failed or found no content, try AI fallbacks
        if not manual_extraction_successful or not result:
            print("Manual extraction failed or found no content. Trying AI fallbacks...")
            
            # Get full document text for AI processing
            document_text = get_full_document_text(url)
            
            if not document_text:
                print("Could not retrieve document text for AI processing")
                return {'tags': [], 'classification': []}
            
            # Try OpenAI first if API key is provided
            if openai_api_key:
                ai_result = extract_tags_with_openai(document_text, openai_api_key)
                if ai_result and (ai_result.get('tags') or ai_result.get('classification')):
                    print("OpenAI extraction successful, using AI results")
                    return ai_result
            
            # Try Bedrock as second fallback if company_bot is provided
            if company_bot:
                ai_result = extract_tags_with_bedrock(document_text, company_bot)
                if ai_result and (ai_result.get('tags') or ai_result.get('classification')):
                    print("Bedrock extraction successful, using AI results")
                    return ai_result
            
            # If all methods fail, return default
            print("All extraction methods failed, returning default tags")
            return {'tags': [], 'classification': []}
        
        return result
        
    except Exception as e:
        print(f"Error extracting sections: {e}")
        return {'tags': [], 'classification': []}

def get_tags_list(url: str, openai_api_key: Optional[str] = None, company_bot=None) -> List[str]:
    """
    Simple function to get just the tags as a clean list with AI fallbacks
    
    Args:
        url: Document URL
        openai_api_key: OpenAI API key for fallback
        company_bot: Company bot object for Bedrock fallback
    
    Returns:
        List of clean tags found in the document
    """
    
    result = extract_tags_from_document_sections(url, openai_api_key, company_bot)
    
    # Combine tags and classification into one list
    all_tags = []
    
    if 'tags' in result:
        all_tags.extend(result['tags'])
    
    if 'classification' in result:
        all_tags.extend(result['classification'])
    
    # Remove duplicates while preserving order
    unique_tags = []
    seen = set()
    for tag in all_tags:
        if tag.lower() not in seen:
            unique_tags.append(tag)
            seen.add(tag.lower())
    
    return unique_tags

def get_classification_list(url: str, openai_api_key: Optional[str] = None, company_bot=None) -> List[str]:
    """
    Get just the classification items as a clean list with AI fallbacks
    
    Args:
        url: Document URL
        openai_api_key: OpenAI API key for fallback
        company_bot: Company bot object for Bedrock fallback
    
    Returns:
        List of clean classification items
    """
    
    result = extract_tags_from_document_sections(url, openai_api_key, company_bot)
    return result.get('classification', [])

# Enhanced one-liner functions with AI fallbacks
def extract_tags(url: str, openai_api_key: Optional[str] = None, company_bot=None) -> List[str]:
    """One-liner to extract clean tags from document with AI fallbacks"""
    return get_tags_list(url, openai_api_key, company_bot)

def extract_classification(url: str, openai_api_key: Optional[str] = None, company_bot=None) -> List[str]:
    """One-liner to extract clean classification from document with AI fallbacks"""
    return get_classification_list(url, openai_api_key, company_bot)

# Example usage with all options
def get_doc_tags_from_ai(file_url, company_bot):
    # # Test 1: Manual extraction only
    # print("1. Manual extraction only:")
    # tags = get_tags_list(file_url)
    # print(f"Tags: {tags}")
    
    # print("\n2. With OpenAI fallback:")
    # # Test 2: With OpenAI fallback
    # tags_with_openai = get_tags_list(file_url, openai_api_key=openai_key)
    # print(f"Tags: {tags_with_openai}")
    
    # Test 3: With Bedrock fallback
    auto_tags = get_tags_list(file_url, company_bot=company_bot)
    print(f"Tags: {auto_tags}")
    return auto_tags
